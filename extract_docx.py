#!/usr/bin/env python3
"""Extract the paired examples and grammar notes from the source DOCX.

The generated JavaScript is intentionally usable from file:// so the learning app
can be opened directly without a server or a build step.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document


SECTION_RE = re.compile(r"^([1-7])[\u3000\s]+(.+)$")
CIRCLED_RE = re.compile(r"^[①②③④⑤⑥⑦⑧]\s*")
JAPANESE_RE = re.compile(r"[\u3040-\u30ff\u3400-\u9fff]")


def text_of(paragraph) -> str:
    return " ".join(paragraph.text.split()).strip()


def is_bold(paragraph) -> bool:
    return any(run.bold for run in paragraph.runs)


def looks_like_translation(text: str) -> bool:
    return bool(JAPANESE_RE.search(text)) and not text.startswith(("★", "「"))


def extract(source: Path) -> dict:
    document = Document(source)
    paragraphs = document.paragraphs
    title = text_of(paragraphs[0])
    subtitle = text_of(paragraphs[1])
    sections: list[dict] = []
    patterns: list[dict] = []
    examples: list[dict] = []
    current_section: dict | None = None
    current_pattern: dict | None = None

    def add_pattern(label: str) -> dict:
        nonlocal current_pattern
        current_pattern = {
            "id": f"p{len(patterns) + 1:02d}",
            "sectionId": current_section["id"],
            "label": CIRCLED_RE.sub("", label).strip(),
            "gloss": "",
            "note": "",
        }
        patterns.append(current_pattern)
        current_section["patternIds"].append(current_pattern["id"])
        return current_pattern

    i = 2
    while i < len(paragraphs):
        paragraph = paragraphs[i]
        text = text_of(paragraph)
        if not text or text == "重要表現の一覧":
            i += 1
            continue

        section_match = SECTION_RE.match(text)
        if section_match:
            number, heading = section_match.groups()
            current_section = {
                "id": f"s{number}",
                "number": int(number),
                "title": heading,
                "patternIds": [],
            }
            sections.append(current_section)
            current_pattern = None
            # Chapter 6 uses its chapter heading as the sole pattern heading.
            if number == "6":
                add_pattern(heading)
            i += 1
            continue

        next_text = text_of(paragraphs[i + 1]) if i + 1 < len(paragraphs) else ""
        pair = (
            current_section is not None
            and is_bold(paragraph)
            and looks_like_translation(next_text)
        )

        if pair:
            if current_pattern is None:
                add_pattern(current_section["title"])
            examples.append(
                {
                    "id": f"e{len(examples) + 1:03d}",
                    "sectionId": current_section["id"],
                    "patternId": current_pattern["id"],
                    "en": text,
                    "ja": next_text,
                }
            )
            i += 2
            continue

        if text.startswith("★"):
            if current_pattern:
                note = text.removeprefix("★").strip()
                current_pattern["note"] = " ".join(
                    part for part in (current_pattern["note"], note) if part
                )
            i += 1
            continue

        if text.startswith("「") and current_pattern:
            current_pattern["gloss"] = text.strip("「」")
            i += 1
            continue

        if current_section is not None and is_bold(paragraph):
            add_pattern(text)

        i += 1

    quick_reference = []
    if document.tables:
        for row in document.tables[0].rows[1:]:
            cells = [" ".join(cell.text.split()).strip() for cell in row.cells]
            if len(cells) >= 2 and all(cells[:2]):
                quick_reference.append({"ja": cells[0], "en": cells[1]})

    for pattern in patterns:
        pattern["exampleCount"] = sum(
            example["patternId"] == pattern["id"] for example in examples
        )
    for section in sections:
        section["exampleCount"] = sum(
            example["sectionId"] == section["id"] for example in examples
        )

    if len(examples) != 66:
        raise ValueError(f"Expected 66 example pairs, extracted {len(examples)}")

    return {
        "title": title,
        "subtitle": subtitle,
        "sourceFile": source.name,
        "sections": sections,
        "patterns": patterns,
        "examples": examples,
        "quickReference": quick_reference,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    payload = extract(args.source)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(
        "window.RELATION_DATA = "
        + json.dumps(payload, ensure_ascii=False, indent=2)
        + ";\n",
        encoding="utf-8",
    )
    print(
        f"Extracted {len(payload['examples'])} examples, "
        f"{len(payload['patterns'])} patterns, and "
        f"{len(payload['sections'])} sections to {args.output}"
    )


if __name__ == "__main__":
    main()
